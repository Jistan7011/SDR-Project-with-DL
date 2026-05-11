from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "LOW_COST_SDR_BLIND_MODULATION_TWO_PAGE_PAPER.docx"


def set_cell_text(cell, text, bold=False, size=7.0):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    for tag, value in (("w:top", "60"), ("w:left", "60"), ("w:bottom", "60"), ("w:right", "60")):
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        elem = tc_mar.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            tc_mar.append(elem)
        elem.set(qn("w:w"), value)
        elem.set(qn("w:type"), "dxa")


def add_para(doc, text, size=9.6, bold=False, align=None, after=3):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.55)
section.bottom_margin = Cm(1.45)
section.left_margin = Cm(1.55)
section.right_margin = Cm(1.55)

styles = doc.styles
styles["Normal"].font.name = "Malgun Gothic"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
styles["Normal"].font.size = Pt(9.6)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(5)
r = title.add_run("저가 SDR 기반 블라인드 BASK/BFSK/BPSK 변조 분류와 BASK 흡수 오류 분석")
r.bold = True
r.font.name = "Malgun Gothic"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
r.font.size = Pt(14)

add_para(doc, "요약", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
add_para(
    doc,
    "본 논문에서는 HackRF One 송신기와 RTL-SDR Blog V4 수신기로 수집한 실제 OTA RF 신호를 대상으로 딥러닝 기반 블라인드 BASK/BFSK/BPSK 변조 분류를 수행하였다. 모델 입력에는 preamble, sync, payload, CRC와 같은 protocol 정보를 사용하지 않았으며, window-level IQ 신호와 물리 기반 feature만 사용하였다. 실험 과정에서 가장 큰 문제는 BFSK와 BPSK가 BASK로 잘못 분류되는 비대칭 혼동 오류였고, 본 논문에서는 이를 BASK absorption error로 정의하였다. 실험 5~9에서 5-channel ResNet baseline, 2-stage classifier, unknown/ambiguous guard, multi-task learning, RF preprocessing을 비교한 결과, 단순 정확도는 실험 5가 0.7308로 가장 높았으나 absorption이 0.1894로 컸다. 실험 8은 absorption을 0.1272까지 낮추었고, 실험 9의 RF-preprocessed ResNet1D는 accuracy 0.7294, worst recall 0.7200, absorption 0.1319로 가장 실용적인 균형을 보였다.",
)
add_para(doc, "주요어: SDR, 딥러닝, 블라인드 변조 분류, BASK, BFSK, BPSK, OTA, RF 신호 분석", size=9.0, bold=True)

add_heading(doc, "1. 서론")
add_para(
    doc,
    "사이버 전자전 및 신호정보 환경에서는 통신 규약을 알 수 없는 송신원의 RF 신호를 빠르게 분석해야 한다. 변조 방식 식별은 복조, 프레임 동기화, payload recovery로 이어지는 첫 단계이므로 blind modulation classification은 unknown RF signal analysis의 핵심 기능이다. 기존 AMC는 higher-order statistics나 cyclostationary feature에 의존했지만, 최근 연구는 raw I/Q 신호를 CNN 또는 ResNet에 입력하여 변조 특징을 자동 학습하는 방향으로 발전하였다[1][3].",
)
add_para(
    doc,
    "O'Shea 등의 OTA 연구는 안정적인 SDR 장비와 대규모 데이터가 있을 때 딥러닝 기반 RF 분류가 높은 성능을 낼 수 있음을 보였고[3], Radio Transformer Network는 frequency offset, phase rotation, timing offset과 같은 동기화 문제가 분류 성능에 직접 영향을 준다는 점을 제시하였다[2]. 본 연구는 이와 달리 HackRF One과 RTL-SDR Blog V4라는 저가 SDR 조합, 제한된 OTA 데이터, session-held-out 평가 조건에서 실험하였다. 따라서 본 논문의 초점은 최고 accuracy뿐 아니라 실제 저가 OTA 환경에서 반복되는 BASK absorption 오류를 분석하고 완화하는 데 있다.",
)

add_heading(doc, "2. 실험 환경 및 데이터 구성")
add_para(
    doc,
    "실험은 433 MHz 대역에서 HackRF One으로 BASK/BFSK/BPSK 신호를 송신하고 RTL-SDR Blog V4로 수신하는 방식으로 수행하였다. 실험 5부터는 실질적인 blind classifier 비교를 위해 2048-sample window와 5-channel feature [I, Q, magnitude, instantaneous_frequency, differential_phase]를 사용하였다. magnitude는 BASK의 envelope 변화, instantaneous frequency는 BFSK의 frequency shift, differential phase는 BPSK의 phase transition을 드러내기 위한 feature이다.",
)
add_para(
    doc,
    "모델 입력에는 frame 구조나 payload 정보를 넣지 않았다. 즉, classifier는 protocol-blind 조건에서 수신 window만 보고 세 변조 중 하나를 예측한다. 실험 4에서 oracle modulation을 사용하면 CRC pass rate가 0.9778까지 상승했지만, classifier 기반 recovery는 약 0.70 수준에 머물렀다. 따라서 이후 실험 5~9에서는 전체 recovery보다 modulation classifier 자체의 성능과 BASK absorption 감소에 집중하였다.",
)

add_heading(doc, "3. 실험 방법")
add_para(
    doc,
    "실험 5는 30-session dataset과 5-channel feature를 사용한 ResNet1D baseline이다. accuracy는 0.7308로 가장 높았지만, BFSK 156개와 BPSK 147개가 BASK로 분류되어 absorption이 0.1894로 나타났다. 실험 6은 Stage 1에서 BASK/non-BASK를 나누고 Stage 2에서 BFSK/BPSK를 분류하는 2-stage classifier를 적용했지만, Stage 1의 hard gate 오류를 복구할 수 없어 absorption이 0.2338로 악화되었다.",
)
add_para(
    doc,
    "실험 7은 unknown-protocol 환경을 고려하여 항상 세 class 중 하나를 강제 선택하지 않고 ambiguous 또는 unknown 상태를 허용하였다. confidence, entropy, envelope evidence, frequency evidence, phase evidence를 함께 사용하여 후보를 보존한 결과 hard-BASK rate는 0.0108로 낮아졌다. 실험 8은 3-class head, BASK/non-BASK head, BFSK/BPSK masked head를 결합한 MultiTaskResNet을 사용하여 absorption을 0.1272까지 줄였다. 실험 9는 RF preprocessing/canonicalization을 적용한 뒤 ResNet1D와 FusionResNet을 비교하였고, 최종적으로 RF-preprocessed ResNet1D seed44를 선택하였다.",
)

add_heading(doc, "4. 실험 결과 및 비교")
table = doc.add_table(rows=1, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
headers = ["실험", "핵심 접근법", "대표 모델", "Accuracy", "Worst Recall", "Absorption"]
for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8.0)
rows = [
    ["Exp5", "5-channel baseline", "ResNet1D", "0.7308", "0.7050", "0.1894"],
    ["Exp6", "2-stage hard gate", "Binary ResNet x2", "0.7238", "0.6512", "0.2338"],
    ["Exp7", "unknown guard", "ResNet + rule", "0.7246", "약 0.7167", "hard-BASK 0.0108"],
    ["Exp8", "multi-task learning", "MultiTaskResNet", "0.7241", "0.7194", "0.1272"],
    ["Exp9", "RF preprocessing", "ResNet1D", "0.7294", "0.7200", "0.1319"],
]
for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        set_cell_text(cells[i], value, size=7.8)

add_para(
    doc,
    "결과적으로 실험 5는 raw accuracy가 가장 높았지만 non-BASK 신호를 BASK로 흡수하는 오류가 컸다. 실험 8은 absorption 감소 효과가 가장 컸고, 실험 9는 accuracy를 실험 5와 거의 같은 수준으로 유지하면서 worst recall과 absorption을 함께 개선하였다. 따라서 최종 실용 모델은 실험 9의 RF-preprocessed ResNet1D로 판단한다. 다만 방법론적으로 가장 의미 있는 개선은 실험 8의 multi-task boundary learning이다.",
)

add_heading(doc, "5. 결론")
add_para(
    doc,
    "본 연구는 저가 SDR 기반 OTA 환경에서 딥러닝을 이용한 blind BASK/BFSK/BPSK 변조 분류를 수행하고, 반복적으로 발생한 BASK absorption 오류를 분석하였다. 단순 정확도 기준으로는 실험 5가 가장 높았으나, 실제 blind RF analysis에서는 absorption과 worst recall이 함께 중요하다. 실험 8은 absorption을 가장 크게 낮추었고, 실험 9는 RF preprocessing과 ResNet1D를 결합하여 가장 균형 잡힌 최종 성능을 보였다. 향후 연구에서는 symbol-aware window selection, multi-window voting, capture-level decision, hard-negative mining, learnable synchronization 구조를 적용하여 session variation과 고신뢰 오분류 문제를 줄일 필요가 있다.",
)

add_heading(doc, "참고문헌")
refs = [
    "[1] T. J. O'Shea, J. Corgan, and T. C. Clancy, Convolutional Radio Modulation Recognition Networks, 2016.",
    "[2] T. J. O'Shea, L. Pemula, D. Batra, and T. C. Clancy, Radio Transformer Networks: Attention Models for Learning to Synchronize in Wireless Systems, 2016.",
    "[3] T. O'Shea, T. Roy, and T. C. Clancy, Over-the-Air Deep Learning Based Radio Signal Classification, 2018.",
    "[4] T. J. O'Shea, N. West, M. Vondal, and T. C. Clancy, Semi-Supervised Radio Signal Identification, 2017.",
    "[5] S. Zhou et al., A Robust Modulation Classification Method Using Convolutional Neural Networks, EURASIP Journal on Advances in Signal Processing, 2019.",
    "[6] H. Han et al., Automatic Modulation Classification Based on Deep Feature Fusion for High Noise Level and Large Dynamic Input, Sensors, 2021.",
    "[7] R. Zhang et al., A Novel Automatic Modulation Classification Method Using Attention Mechanism and Hybrid Parallel Neural Network, 2021.",
]
for ref in refs:
    add_para(doc, ref, size=8.2, after=0)

doc.save(OUT)
print(OUT)
